"""One-off script to generate docs/AflAlert_Technical_Report.docx.

Mirrors the section structure/styling of the CSC1304 VendorSync report
template: cover page, contents, grey section-header banners, numbered
subsections, appendices. Not part of the app build -- run manually with
`python scripts/generate_report.py` whenever the report content needs
regenerating from this source.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# AflAlert brand palette (lib/constants/app_colors.dart) rather than the
# VendorSync template's generic grey/red, so the report reads as this app's
# own document.
BANNER_FILL = "00462D"  # AppColors.primary
BANNER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
HEADING_COLOR = RGBColor(0x00, 0x46, 0x2D)  # AppColors.primary
PLACEHOLDER = RGBColor(0xC6, 0x28, 0x28)  # AppColors.error

# Contributions summarised from `git log --author=<email>` commit history and
# per-author touched-file counts, not self-reported -- reflects what each
# person's commits actually show, not a claimed division of labour.
TEAM = [
    (
        "NABACHWA PRISCILLA LUTAYA",
        "25/U/0350",
        [
            "Configured Firebase and platform build setup (Android/iOS/Windows); "
            "resolved recurring build/run issues (Gradle, Kotlin)",
            "Built and iterated the home, history, and analysis screens",
            "Implemented image storage handling and maize image resizing for scans",
            "Worked on the forgot-password flow, settings screen, and legal section",
            "Contributed to localization of the login page and general UI",
        ],
    ),
    (
        "NALUYIMA LAUREEN",
        "25/U/03511/PS",
        [
            "Worked on iOS and Android platform configuration",
            "Built the login screen and created the forgot-password screen",
            "Added Google sign-in branding assets (logo sizing/placement)",
            "Contributed to the analysis and settings screens",
            "Added Luganda-language localization changes",
            "Resolved recurring merge conflicts during branch integration",
        ],
    ),
    (
        "ATWINE SHEILA BAGUMIRE",
        "25/U/03348/EVE",
        [
            "Built out the home screen extensively, plus registration, login, and "
            "settings screens",
            "Implemented the history screen and alert notifications (weather-based "
            "alerts, daily tips)",
            "Added the voice assistant integration and its icon",
            "Wrote Terms of Service and legal/guideline content",
            "Localized major screens into English and Luganda (app_localizations)",
            "Led resolution of merge conflicts across the application branch "
            "integration",
        ],
    ),
    (
        "NSUBUGA BRIAN",
        "25/U/03546/PSA",
        [
            "Set up localization support (flutter_localizations, intl "
            "dependencies) and initialized the Luganda ARB file",
            "Restructured the l10n folder and generated the English/Luganda "
            "localization delegate files",
            "Added the history screen",
            "Edited the profile screen and project dependencies (pubspec.yaml)",
            "Worked on platform-specific TFLite and Firebase Storage services",
            "Resolved merge conflicts by reconciling incoming changes",
        ],
    ),
    (
        "KIZITO JUMA",
        "25/U/26620",
        [
            "Produced the Luganda translation file (app_lg.arb) and localization "
            "strings",
            "Built the notifications screen and result screen",
            "Worked on the welcome, login, and registration pages",
            "Sized and integrated the AflAlert logo across screens",
            "Integrated the TFLite/ML classification service and weather service",
            "Resolved merge conflicts on the home screen",
        ],
    ),
]

doc = Document()

# ---- base style ----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)


def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def add_banner(text, toc_level=1):
    # A shaded paragraph (not a table) so this is a real Heading-styled
    # paragraph in the main body flow -- Word's native TOC field only
    # reliably discovers headings sitting directly in the body, not ones
    # nested inside table cells.
    p = doc.add_paragraph()
    if toc_level is not None:
        p.style = doc.styles[f"Heading {toc_level}"]
    shade_paragraph(p, BANNER_FILL)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BANNER_TEXT
    doc.add_paragraph()


def add_toc():
    """Insert a real Word TOC field (auto-updating, not a static list)."""
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to load the table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    for el in (fld_begin, instr, fld_separate, placeholder, fld_end):
        r.append(el)

    # Ask Word to update all fields (including this TOC) automatically when
    # the document is opened, so page numbers are correct without the user
    # needing to know about "Update Field".
    settings_el = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_el.append(update_fields)


def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR


def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_placeholder_para(label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.italic = True
    run.font.color.rgb = PLACEHOLDER
    return p


def add_bullets(items, bold_lead=None):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if bold_lead and item.startswith(bold_lead[0]):
            pass
        p.add_run(item)


def add_lettered(items):
    letters = "abcdefghijklmnopqrstuvwxyz"
    for i, (lead, rest) in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        r1 = p.add_run(f"{letters[i]}) ")
        r1.bold = True
        if lead:
            r2 = p.add_run(f"{lead}: ")
            r2.bold = True
        p.add_run(rest)


# ---- Cover page ------------------------------------------------------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = title_p.add_run("Technical Report")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = HEADING_COLOR

for_p = doc.add_paragraph()
for_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = for_p.add_run("for")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x70, 0x79, 0x73)  # AppColors.grey

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = name_p.add_run("AflAlert")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = HEADING_COLOR

for _ in range(6):
    doc.add_paragraph()

prep_p = doc.add_paragraph()
prep_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = prep_p.add_run("Prepared by")
r.bold = True

doc.add_paragraph()

group_p = doc.add_paragraph()
group_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = group_p.add_run("Group Name: ")
r.bold = True
r2 = group_p.add_run("GROUP 4")

table = doc.add_table(rows=len(TEAM), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row, (name, reg_no, _contributions) in zip(table.rows, TEAM):
    for ci, col_text in enumerate([reg_no, name]):
        cell = row.cells[ci]
        p = cell.paragraphs[0]
        r = p.add_run(col_text)
        r.font.size = Pt(10)

doc.add_paragraph()

for label, value, is_placeholder in [
    ("Mentor:", "MR. ASIIMWE PADDY", False),
    ("Course:", "CSC 1304 Practical Skills Development", False),
    ("Date:", "1 August 2026", False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r2 = p.add_run(value)
    if is_placeholder:
        r2.italic = True
        r2.font.color.rgb = PLACEHOLDER

doc.add_page_break()

# ---- Contents ----------------------------------------------------------
add_banner("Contents", toc_level=None)
add_toc()

doc.add_page_break()

# ---- Abstract ---------------------------------------------------------------
add_banner("Abstract")
add_para(
    "AflAlert is a mobile-first, AI-assisted risk-detection application built to help maize "
    "farmers and grain handlers identify aflatoxin contamination risk without needing "
    "laboratory access. It combines two complementary detection tiers: an on-device image-"
    "classification model that flags visible mould/spoilage risk from a photo of maize "
    "kernels, and a chemical lateral-flow test-strip reader that estimates a contamination "
    "level in parts per billion (ppb) from a photographed, reacted strip."
)
add_para(
    "Every scan is paired with guided next-step recommendations sourced from UNBS and MAAIF "
    "standards, alongside proactive rain/weather alerts that help farmers protect drying "
    "grain before it is exposed to moisture — a primary driver of aflatoxin growth."
)
add_para(
    "AflAlert also provides scan history, exportable PDF reports, a hands-free voice "
    "assistant, and multilingual support (English and Luganda), giving smallholder farmers "
    "and grain handlers a practical, accessible tool for reducing contamination risk and "
    "protecting both health and income."
)

doc.add_page_break()

# ---- 1 Introduction -----------------------------------------------------
add_banner("1  Introduction")
add_para(
    "This document presents the technical report for AflAlert, an aflatoxin risk-detection "
    "application developed as part of [Course Code and Title]. It provides an overview of "
    "the project's objectives, system design, functionality, implementation outcomes, "
    "limitations, and planned future enhancements."
)
add_para(
    "The report serves as a formal record of the project's development process and "
    "achievements, detailing how AflAlert addresses the lack of accessible aflatoxin "
    "screening for farmers through on-device kernel scanning, chemical test-strip reading, "
    "guided recommendations, and proactive weather alerts."
)

add_heading("1.1  User Challenge", level=2)
add_para(
    "Aflatoxin contamination in maize is a major food-safety and economic risk in many "
    "farming communities, but affected farmers and grain handlers rarely have access to "
    "laboratory testing. Visual inspection alone is unreliable — early or low-level "
    "contamination isn't visible to the eye — while commercial lab testing is costly, slow, "
    "and often geographically inaccessible for smallholder and rural sellers."
)
add_para(
    "Meanwhile, unpredictable rain during the drying period is one of the most common causes "
    "of moisture re-uptake that drives mould and toxin growth; without timely local weather "
    "warnings, farmers often cannot cover or move drying grain in time. Grain handlers and "
    "farmers therefore need a fast, low-cost, easy first-line way to assess risk directly at "
    "the point of harvest or sale, along with practical guidance on what to do about a given "
    "result, and warning of conditions — like incoming rain — that raise their risk."
)

add_heading("1.2  Project Goals", level=2)
add_para(
    "AflAlert's goal is to put a first-line aflatoxin risk check directly into the hands of "
    "farmers and grain handlers, using only a smartphone camera — no lab, no reagents beyond "
    "a commercially available test strip, no specialist training required."
)
add_para(
    "It combines two complementary scanning tiers (kernel image analysis and strip-based "
    "chemical reading) so a user can get an assessment even without a strip on hand, and a "
    "more chemically-grounded reading when one is used. Every result is paired with "
    "actionable guidance rather than a bare number, and the app proactively protects against "
    "a leading cause of contamination — rain exposure during drying — through location-based "
    "weather alerts. Voice assistance, scan history, and shareable PDF reports round out the "
    "experience so results can be acted on, tracked over time, and shared with buyers, "
    "extension officers, or family members who help with drying and storage decisions."
)

add_heading("1.3  Functional Requirements", level=2)

add_heading("1.3.1.  Kernel Image Scanning", level=3)
add_lettered([
    ("Capture or select a photo", "Allow the user to photograph a sample of maize kernels via the in-app camera or pick an existing photo from the gallery."),
    ("On-device classification", "Run the captured image through a bundled TensorFlow Lite model to classify the sample (e.g. healthy vs. mouldy/spoiled) and flag likely non-maize photos."),
    ("Confidence and rejection handling", "Reject low-confidence or clearly non-maize images (colour-profile mismatch, model rejection, low confidence) rather than returning a misleading result."),
])

add_heading("1.3.2.  Test Strip Scanning", level=3)
add_lettered([
    ("Guided capture", "Present a capture frame that guides the user to align a reacted lateral-flow test strip (Control line near the top, Test line near the bottom) and checks light/focus quality before use."),
    ("Strip plausibility check", "Reject photos that don't plausibly resemble a strip cassette (e.g. not sufficiently pale/low-saturation) before attempting line detection."),
    ("Line detection", "Locate the Control (C) and Test (T) line bands from a row-by-row luminance profile of the image, using a local-background baseline to stay robust to uneven lighting."),
    ("Optical density and ppb calculation", "Compute the optical density of each line relative to its local background, derive a T/C ratio, and translate that ratio into an estimated contamination reading in ppb."),
    ("Validity (QC) checks", "Flag a scan as invalid if no strip/lines are detected, or if the Control line isn't dark enough to indicate a properly developed run, rather than presenting a misleading low-ppb result."),
])

add_heading("1.3.3.  Guided Recommendations", level=3)
add_lettered([
    ("Result-linked guidance", "Attach next-step guidance to every scan result (kernel or strip), sourced from UNBS and MAAIF standards, based on whether the result is within or above the safe threshold."),
])

add_heading("1.3.4.  Weather and Rain Alerts", level=3)
add_lettered([
    ("Daily morning summary", "Provide a once-daily weather summary based on the user's location."),
    ("Urgent rain alerts", "Independently monitor short-range forecasts and push an urgent alert when rain is imminent, so drying grain can be covered or moved in time, including while the app is closed (Android background scheduling)."),
])

add_heading("1.3.5.  Voice Assistant", level=3)
add_lettered([
    ("Hands-free interaction", "Allow the user to trigger scans and ask questions about results or recommendations using speech input and spoken responses."),
    ("Available across the app", "Surface the assistant from the home, history, notifications, and results screens (kernel and strip), not just a single entry point, so it's reachable wherever a user might want it."),
    ("Spoken result summaries", "When a scan is triggered from the assistant, speak a summary of the result (safe/unsafe classification) aloud via text-to-speech as soon as it's ready, using the same label logic as the on-screen result so the two never disagree."),
])

add_heading("1.3.6.  History, Reports and Accounts", level=3)
add_lettered([
    ("Scan history", "Store and display a history of past kernel and strip scans per authenticated user."),
    ("PDF report export", "Generate and share/download a PDF report for a given scan result."),
    ("Authentication", "Support email/password and Google sign-in, OTP verification, password reset, and biometric app unlock."),
    ("Localization", "Support English and Luganda throughout the app."),
])

doc.add_page_break()

# ---- 2 Project Results ---------------------------------------------------
add_banner("2  Project Results")

add_heading("2.1  Product Design", level=2)
add_para(
    "AflAlert is a cross-platform mobile application built with Flutter, using Firebase for "
    "authentication (email/password, Google Sign-In, OTP), real-time data storage (Cloud "
    "Firestore), file storage (Firebase Storage), server-side logic (Cloud Functions), and "
    "app integrity verification (Firebase App Check)."
)
add_para(
    "On-device machine learning is handled via tflite_flutter, running a bundled "
    "classification model for kernel scanning, while the test-strip reader uses "
    "deterministic pixel/image processing (the image package) rather than a trained model, "
    "since the T/C line-reading problem is well suited to direct optical-density calculation."
)
add_para(
    "Background scheduling for rain/weather alerts uses WorkManager (Android) so alerts can "
    "fire even when the app isn't open. The codebase separates concerns cleanly into screens "
    "(UI), services (business logic — classification, strip analysis, Firebase access, "
    "alerts, PDF generation, voice), and localization resources, making each detection tier "
    "and supporting feature independently maintainable and testable."
)
add_placeholder_para("[Insert system architecture diagram here.]")

add_heading("2.2  Product Functionality and Screenshots", level=2)

add_para("a) Kernel Scan Functionality", bold=True)
add_para(
    "Users capture or select a photo of maize kernels; the app classifies the sample "
    "on-device and returns a risk result with guided recommendations."
)
add_placeholder_para("[Insert kernel scan screenshots here.]")

add_para("b) Test Strip Scan Functionality", bold=True)
add_para(
    "Users perform the physical extraction/strip test, then photograph the reacted strip "
    "through a guided capture flow; the app reads the Control and Test lines, computes an "
    "optical-density ratio, and returns an estimated ppb reading with a safe/unsafe "
    "classification and guidance."
)
add_placeholder_para("[Insert strip scan screenshots here.]")

add_para("c) Alerts, History and Assistant Functionality", bold=True)
add_para(
    "Users receive daily weather summaries and urgent rain alerts, review past scans and "
    "download PDF reports, and can interact with a voice assistant for hands-free scanning "
    "and Q&A. The assistant is reachable from the home, history, notifications, and results "
    "screens, and speaks a result summary aloud when a scan it triggered finishes."
)
add_placeholder_para("[Insert alerts/history/voice-assistant screenshots here.]")

add_heading("2.3  Project Website and Repository", level=2)
for label, value, is_placeholder in [
    ("GitHub Repository Link", "https://github.com/priscillanabachwa/AFLALERT", False),
    ("Project Website Link", "[Project website URL]", True),
    ("Running System Link", "[Deployed app / demo URL]", True),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r2 = p.add_run(value)
    if is_placeholder:
        r2.italic = True
        r2.font.color.rgb = PLACEHOLDER

doc.add_page_break()

# ---- 3 Limitations and Next Steps ----------------------------------------
add_banner("3  Limitations and Next Steps")

add_heading("3.1  Limitations", level=2)
add_lettered([
    ("Placeholder ppb calibration", "The strip reader's ppb figure is currently derived from a monotonic placeholder curve rather than a lab-fitted calibration, so it should be treated as a relative/qualitative signal rather than a certified quantitative reading."),
    ("Gallery-picked strip photos are less reliable", "Photos selected from the gallery (rather than the guided in-app camera capture) aren't cropped/oriented to the expected frame, making line detection less reliable on those images."),
    ("Background alerts are Android-only", "iOS has no platform guarantee for background task timing, so the proactive rain/weather alert currently only runs reliably on Android."),
    ("No confirmatory lab pathway", "The app offers a first-line risk read, not a certified lab-equivalent result, and does not yet route high-risk results to confirmatory testing."),
])

add_heading("3.2  Next Steps", level=2)
add_lettered([
    ("Calibrate the strip reader", "Run a dilution series of certified-reference aflatoxin standards through the strip and app pipeline, and fit a proper dose-response curve (e.g. 4-parameter logistic) to replace the placeholder ppb formula."),
    ("Improve gallery-photo robustness", "Add stricter structural validation for non-guided (gallery) strip photos, or restrict strip scanning to in-app camera capture only."),
    ("iOS background alerts", "Investigate iOS-appropriate scheduling (e.g. BGTaskScheduler) to bring rain alerts to iOS."),
    ("Confirmatory testing pathway", "Add a recommended next step to route high-risk results toward certified lab confirmation."),
])

doc.add_page_break()

# ---- References ------------------------------------------------------------
add_banner("References")
refs = [
    "[1] Google, \"Flutter documentation,\" Flutter.dev. [Online]. Available: "
    "https://flutter.dev/docs.",
    "[2] Google, \"Firebase documentation,\" Firebase.google.com. [Online]. Available: "
    "https://firebase.google.com/docs.",
    "[3] Dart team, \"Dart language guides,\" Dart.dev. [Online]. Available: "
    "https://dart.dev/guides.",
    "[4] Google, \"TensorFlow Lite,\" TensorFlow.org. [Online]. Available: "
    "https://www.tensorflow.org/lite.",
    "[5] \"image,\" Dart package, pub.dev. [Online]. Available: "
    "https://pub.dev/packages/image.",
    "[6] Codex Alimentarius Commission, \"General standard for contaminants and "
    "toxins in food and feed,\" CODEX STAN 193-1995, FAO/WHO, Rome, Italy, "
    "1995 (rev. 2019) — sets the international reference maximum levels for "
    "aflatoxins in food that UNBS/MAAIF grain guidance is aligned with.",
    "[7] Uganda National Bureau of Standards (UNBS), \"Maize grains — "
    "Specification,\" Kampala, Uganda. [Online]. Available: https://www.unbs.go.ug.",
    "[8] Ministry of Agriculture, Animal Industry and Fisheries (MAAIF), "
    "\"Post-harvest handling and grain storage guidelines,\" Kampala, Uganda. "
    "[Online]. Available: https://www.agriculture.go.ug.",
    "[9] \"flutter_tts,\" Dart package, pub.dev. [Online]. Available: "
    "https://pub.dev/packages/flutter_tts.",
    "[10] \"workmanager,\" Dart package, pub.dev. [Online]. Available: "
    "https://pub.dev/packages/workmanager.",
]
for r in refs:
    doc.add_paragraph(r)

doc.add_page_break()

# ---- Appendix A ------------------------------------------------------------
add_banner("Appendix A – Project Work Plan")

# Usable content width = page width minus margins.
content_width = section.page_width - section.left_margin - section.right_margin

pic_p = doc.add_paragraph()
pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_p.add_run().add_picture("docs/assets/gantt_chart_overview.png", width=content_width)

doc.add_page_break()

pic_p2 = doc.add_paragraph()
pic_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_p2.add_run().add_picture("docs/assets/gantt_chart_phase_details.png", width=content_width)

doc.add_page_break()

# ---- Appendix B ------------------------------------------------------------
add_banner("Appendix B – Contribution by Team Members")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "No."
hdr[1].text = "Team Member"
hdr[2].text = "Contribution"
for c in hdr:
    shade_cell(c, BANNER_FILL)
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = BANNER_TEXT

for i, (name, _reg_no, contributions) in enumerate(TEAM, start=1):
    row = table.add_row().cells
    row[0].text = str(i)
    p = row[1].paragraphs[0]
    p.add_run(name)
    cell = row[2]
    first = True
    for line in contributions:
        p2 = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p2.add_run(f"✓ {line}")

doc.save("docs/AflAlert_Technical_Report.docx")
print("Saved docs/AflAlert_Technical_Report.docx")
