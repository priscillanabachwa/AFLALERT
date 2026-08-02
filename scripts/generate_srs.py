"""One-off script to generate the AflAlert Software Requirements Specification.

Mirrors the section structure/styling of generate_report.py (same banner
colours, cover-page layout, helper functions) but follows the CSC1304
SRS-template-CSC1304 structure (Introduction, Overall Description, Specific
Requirements, Non-functional Requirements, Other Requirements, Group Log)
instead of the technical-report structure. Not part of the app build --
run manually with `python scripts/generate_srs.py` whenever the SRS content
needs regenerating from this source. Saves the finished document straight to
the Desktop rather than docs/, since that's where this is submitted from.
"""

import os
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Same AflAlert brand palette as generate_report.py (lib/constants/app_colors.dart)
BANNER_FILL = "00462D"  # AppColors.primary
BANNER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
HEADING_COLOR = RGBColor(0x00, 0x46, 0x2D)  # AppColors.primary
PLACEHOLDER = RGBColor(0xC6, 0x28, 0x28)  # AppColors.error

TEAM = [
    ("NABACHWA PRISCILLA LUTAYA", "25/U/0350", "priscillanabachwa4@gmail.com"),
    ("NALUYIMA LAUREEN", "25/U/03511/PS", None),
    ("ATWINE SHEILA BAGUMIRE", "25/U/03348/EVE", None),
    ("NSUBUGA BRIAN", "25/U/03546/PSA", None),
    ("KIZITO JUMA", "25/U/26620", None),
]

DESKTOP_DIR = os.path.join(os.path.expanduser("~"), "Desktop")
OUTPUT_PATH = os.path.join(DESKTOP_DIR, "AflAlert_SRS.docx")

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS_DIR = os.path.join(REPO_ROOT, "docs", "assets")
USE_CASE_DIAGRAM_PATH = os.path.join(ASSETS_DIR, "use_case_diagram.png")
ARCH_DIAGRAM_PATH = os.path.join(ASSETS_DIR, "architecture_diagram.png")


# ---------------------------------------------------------------------------
# Use case diagram (generated with Pillow -- no diagramming lib available)
# ---------------------------------------------------------------------------
def build_use_case_diagram(path):
    W, H = 1700, 1150
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)

    def load_font(name, size):
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            return ImageFont.load_default()

    font_title = load_font("arialbd.ttf", 28)
    font_actor = load_font("arialbd.ttf", 19)
    font_uc = load_font("arial.ttf", 17)

    primary = (0x00, 0x46, 0x2D)
    line_color = (0x8A, 0x8A, 0x8A)
    black = (0x20, 0x20, 0x20)

    def text_center(xy, text, font, fill, anchor="mm"):
        d.text(xy, text, font=font, fill=fill, anchor=anchor)

    def wrap_text(text, font, max_width):
        words = text.split()
        lines, cur = [], ""
        for w in words:
            trial = (cur + " " + w).strip()
            bbox = d.textbbox((0, 0), trial, font=font)
            if bbox[2] - bbox[0] <= max_width or not cur:
                cur = trial
            else:
                lines.append(cur)
                cur = w
        if cur:
            lines.append(cur)
        return lines

    def stick_figure(cx, top_y, color, label):
        r = 20
        d.ellipse([cx - r, top_y, cx + r, top_y + 2 * r], outline=color, width=4)
        body_top = top_y + 2 * r
        body_bottom = body_top + 60
        d.line([cx, body_top, cx, body_bottom], fill=color, width=4)
        d.line([cx - 32, body_top + 22, cx + 32, body_top + 22], fill=color, width=4)
        d.line([cx, body_bottom, cx - 28, body_bottom + 48], fill=color, width=4)
        d.line([cx, body_bottom, cx + 28, body_bottom + 48], fill=color, width=4)
        label_y = body_bottom + 60
        for i, ln in enumerate(wrap_text(label, font_actor, 190)):
            text_center((cx, label_y + i * 24), ln, font_actor, black)
        return (cx, top_y, body_bottom)

    # Title
    text_center((W // 2, 30), "AflAlert — Use Case Diagram", font_title, primary)

    # Actors
    actor1 = stick_figure(130, 330, black, "Farmer / Grain Handler")
    actor2 = stick_figure(130, 800, black, "Background Scheduler (WorkManager)")

    # System boundary
    box_x0, box_y0, box_x1, box_y1 = 380, 70, 1660, 1100
    d.rounded_rectangle([box_x0, box_y0, box_x1, box_y1], radius=18, outline=primary, width=4)
    text_center(((box_x0 + box_x1) // 2, box_y0 + 30), "AflAlert System", font_title, primary)

    usecases = [
        "Register / Sign In",
        "Scan Maize Kernels (Tier 1)",
        "Scan Test Strip (Tier 2)",
        "View Guided Recommendations",
        "Receive Weather & Rain Alerts",
        "Use Voice Assistant",
        "View Scan History",
        "Export PDF Report",
        "Manage Settings & Profile",
    ]

    cols, rows = 3, 3
    margin = 50
    top_pad = 80
    grid_w = (box_x1 - box_x0) - 2 * margin
    grid_h = (box_y1 - box_y0) - top_pad - margin
    cell_w = grid_w / cols
    cell_h = grid_h / rows
    ell_w, ell_h = cell_w - 50, 92

    centers = []
    for idx, label in enumerate(usecases):
        row, col = divmod(idx, cols)
        cx = box_x0 + margin + col * cell_w + cell_w / 2
        cy = box_y0 + top_pad + row * cell_h + cell_h / 2
        centers.append((cx, cy, label))

    # Connector lines first (so ellipses draw on top)
    alert_idx = usecases.index("Receive Weather & Rain Alerts")
    for idx, (cx, cy, label) in enumerate(centers):
        left_edge = (cx - ell_w / 2, cy)
        d.line([actor1[0] + 28, actor1[2] - 10, left_edge[0], left_edge[1]], fill=line_color, width=2)
        if idx == alert_idx:
            d.line([actor2[0] + 28, actor2[2] - 10, left_edge[0], left_edge[1]], fill=line_color, width=2)

    for cx, cy, label in centers:
        d.ellipse([cx - ell_w / 2, cy - ell_h / 2, cx + ell_w / 2, cy + ell_h / 2],
                  outline=primary, width=3, fill="white")
        lines = wrap_text(label, font_uc, ell_w - 30)
        start_y = cy - (len(lines) - 1) * 11
        for i, ln in enumerate(lines):
            text_center((cx, start_y + i * 22), ln, font_uc, black)

    os.makedirs(os.path.dirname(path), exist_ok=True)
    img.save(path)


# ---------------------------------------------------------------------------
# Document build
# ---------------------------------------------------------------------------
build_use_case_diagram(USE_CASE_DIAGRAM_PATH)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
content_width = section.page_width - section.left_margin - section.right_margin


def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def shade_paragraph(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def add_banner(text, toc_level=1):
    p = doc.add_paragraph()
    if toc_level is not None:
        p.style = doc.styles[f"Heading {toc_level}"]
    shade_paragraph(p, BANNER_FILL)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = BANNER_TEXT
    doc.add_paragraph()


def add_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to load the table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    for el in (fld_begin, instr, fld_separate, placeholder, fld_end):
        r.append(el)

    settings_el = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings_el.append(update_fields)


def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = HEADING_COLOR


def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_placeholder_para(label):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.italic = True
    run.font.color.rgb = PLACEHOLDER
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_lettered(items):
    letters = "abcdefghijklmnopqrstuvwxyz"
    for i, (lead, rest) in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        r1 = p.add_run(f"{letters[i]}) ")
        r1.bold = True
        if lead:
            r2 = p.add_run(f"{lead}: ")
            r2.bold = True
        p.add_run(rest)


def add_req_table(rows, headers=("ID", "Requirement")):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], BANNER_FILL)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = BANNER_TEXT
    for row_vals in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_vals):
            row[i].text = val
    doc.add_paragraph()


# ---- Cover page -------------------------------------------------------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = title_p.add_run("Software Requirements")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = HEADING_COLOR

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = title_p2.add_run("Specification")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = HEADING_COLOR

for_p = doc.add_paragraph()
for_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = for_p.add_run("for")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x70, 0x79, 0x73)

name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = name_p.add_run("AflAlert")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = HEADING_COLOR

version_p = doc.add_paragraph()
version_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = version_p.add_run("Version 1.0")
r.font.size = Pt(13)

for _ in range(5):
    doc.add_paragraph()

prep_p = doc.add_paragraph()
prep_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = prep_p.add_run("Prepared by")
r.bold = True

doc.add_paragraph()

group_p = doc.add_paragraph()
group_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = group_p.add_run("Group Name: ")
r.bold = True
group_p.add_run("GROUP 4")

table = doc.add_table(rows=len(TEAM), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row, (name, reg_no, email) in zip(table.rows, TEAM):
    for ci, col_text in enumerate([name, reg_no, email or "<e-mail>"]):
        cell = row.cells[ci]
        p = cell.paragraphs[0]
        r = p.add_run(col_text)
        r.font.size = Pt(10)
        if email is None and ci == 2:
            r.italic = True
            r.font.color.rgb = PLACEHOLDER

doc.add_paragraph()

for label, value, is_placeholder in [
    ("Mentor:", "MR. ASIIMWE PADDY", False),
    ("Course:", "CSC 1304 Practical Skills Development", False),
    ("Date:", "1 August 2026", False),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r2 = p.add_run(value)
    if is_placeholder:
        r2.italic = True
        r2.font.color.rgb = PLACEHOLDER

doc.add_page_break()

# ---- Revisions ---------------------------------------------------------
add_banner("Revisions", toc_level=None)
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, h in enumerate(["Version", "Primary Author(s)", "Description of Version", "Date Completed"]):
    hdr[i].text = h
    shade_cell(hdr[i], BANNER_FILL)
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = BANNER_TEXT
row = table.add_row().cells
row[0].text = "1.0"
row[1].text = "Group 4"
row[2].text = ("Initial SRS draft covering introduction, overall description, "
               "specific requirements, and non-functional requirements.")
row[3].text = "1 August 2026"

doc.add_page_break()

# ---- Contents ------------------------------------------------------------
add_banner("Contents", toc_level=None)
add_toc()
doc.add_page_break()

# ---- 1 Introduction --------------------------------------------------------
add_banner("1  Introduction")
add_para(
    "This document is the Software Requirements Specification (SRS) for AflAlert, a "
    "mobile application that gives maize farmers and grain handlers a first-line "
    "aflatoxin contamination risk check using only a smartphone camera. It was prepared "
    "by Group 4 as part of CSC 1304 Practical Skills Development."
)
add_para(
    "The remainder of this section defines the document's purpose and scope, its "
    "intended readers, the terminology used throughout, and the conventions and "
    "references that apply to the rest of the document."
)

add_heading("1.1  Document Purpose", level=2)
add_para(
    "This SRS specifies the functional and non-functional requirements for AflAlert "
    "version 1.0. It covers the full application: on-device maize kernel scanning, "
    "lateral-flow test-strip reading, guided recommendations, weather/rain alerts, the "
    "voice assistant, scan history and PDF reporting, and account management. It is "
    "intended to be the shared reference that Group 4 designs, implements, and tests "
    "against for the duration of the project, and that the course mentor and "
    "facilitator use to assess whether the delivered app meets its stated requirements."
)

add_heading("1.2  Product Scope", level=2)
add_para(
    "AflAlert lets a farmer or grain handler photograph a sample of maize kernels, or a "
    "reacted lateral-flow test strip after an on-site extraction, and receive an "
    "instant, on-device risk read together with guided next steps sourced from UNBS and "
    "MAAIF standards. It is designed for the point of harvest or sale, where laboratory "
    "testing is rarely available, too slow, or too costly."
)
add_para(
    "Beyond the two scanning tiers, AflAlert helps prevent contamination before it "
    "happens: proactive rain and weather alerts warn farmers in time to cover or move "
    "drying grain, which is one of the most common triggers of mould growth. Scan "
    "history, shareable PDF reports, a hands-free voice assistant, and English/Luganda "
    "localization make the app usable and useful for smallholder farmers, grain "
    "traders, and the extension officers and buyers they share results with. The "
    "product contributes to SDG 2 (Zero Hunger) and SDG 3 (Good Health and Well-Being) "
    "by improving food safety at the grassroots level."
)

add_heading("1.3  Intended Audience and Document Overview", level=2)
add_para(
    "This document is written for three audiences: the Group 4 development team, who "
    "use it as the working reference for what to build and test; the course mentor "
    "(Mr. Asiimwe Paddy) and facilitator, who use it to evaluate whether the project "
    "meets its stated requirements; and, indirectly, the farmers, grain handlers, and "
    "extension officers the product is designed for, whose needs are reflected in "
    "Sections 2 and 3."
)
add_para(
    "Section 1 (this section) introduces the document. Section 2 gives an overall "
    "description of the product, its users, and the environment it runs in. Section 3 "
    "lists specific interface and functional requirements, plus the use-case view. "
    "Section 4 covers non-functional requirements (performance, safety/security, "
    "quality attributes). Section 5 covers other requirements not captured elsewhere. "
    "Readers new to the project should read Sections 1 and 2 first; developers "
    "implementing a specific feature should go straight to the relevant subsection of "
    "Section 3."
)

add_heading("1.4  Definitions, Acronyms and Abbreviations", level=2)
definitions = [
    ("Aflatoxin", "A toxic compound produced by Aspergillus flavus and Aspergillus parasiticus moulds; classified as a Group 1 human carcinogen by IARC."),
    ("App Check", "Firebase App Check — verifies that backend requests originate from the genuine AflAlert app."),
    ("FR", "Functional Requirement."),
    ("MAAIF", "Ministry of Agriculture, Animal Industry and Fisheries (Uganda)."),
    ("NFR", "Non-Functional Requirement."),
    ("OTP", "One-Time Password, used during registration and password reset."),
    ("ppb", "Parts per billion — the unit used for the estimated aflatoxin contamination reading."),
    ("SRS", "Software Requirements Specification (this document)."),
    ("Tier 1", "The kernel-image scanning detection tier."),
    ("Tier 2", "The lateral-flow test-strip scanning detection tier."),
    ("TFLite", "TensorFlow Lite — the on-device machine-learning runtime used for kernel classification."),
    ("UI", "User Interface."),
    ("UNBS", "Uganda National Bureau of Standards."),
    ("WHO", "World Health Organization."),
]
for term, defn in definitions:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{term}: ")
    r1.bold = True
    p.add_run(defn)

add_heading("1.5  Document Conventions", level=2)
add_para(
    "This document follows the CSC 1304 SRS template, itself adapted from IEEE "
    "requirements-document practice. Section and subsection headings follow the "
    "template's numbering. Functional requirements are labelled FR-<area>-<n> and "
    "grouped by functional area (Section 3.2); non-functional requirements are "
    "referenced by section (e.g. §4.1 for performance). Text in "
)
run = doc.paragraphs[-1].add_run("red italics")
run.italic = True
run.font.color.rgb = PLACEHOLDER
doc.paragraphs[-1].add_run(
    " marks information that is not yet finalised (e.g. a teammate's e-mail, or "
    "meeting minutes still to be logged) and must be filled in by the group before "
    "final submission."
)

add_heading("1.6  References and Acknowledgments", level=2)
refs = [
    "[1] Google, \"Flutter documentation,\" Flutter.dev. [Online]. Available: https://flutter.dev/docs.",
    "[2] Google, \"Firebase documentation,\" Firebase.google.com. [Online]. Available: https://firebase.google.com/docs.",
    "[3] Dart team, \"Dart language guides,\" Dart.dev. [Online]. Available: https://dart.dev/guides.",
    "[4] Google, \"TensorFlow Lite,\" TensorFlow.org. [Online]. Available: https://www.tensorflow.org/lite.",
    "[5] \"image,\" Dart package, pub.dev. [Online]. Available: https://pub.dev/packages/image.",
    "[6] Open-Meteo, \"Weather Forecast API,\" open-meteo.com. [Online]. Available: https://open-meteo.com/en/docs.",
    "[7] Codex Alimentarius Commission, \"General standard for contaminants and toxins in food and feed,\" CODEX STAN 193-1995, FAO/WHO, Rome, Italy, 1995 (rev. 2019).",
    "[8] Uganda National Bureau of Standards (UNBS), \"Maize grains — Specification,\" Kampala, Uganda. [Online]. Available: https://www.unbs.go.ug.",
    "[9] Ministry of Agriculture, Animal Industry and Fisheries (MAAIF), \"Post-harvest handling and grain storage guidelines,\" Kampala, Uganda. [Online]. Available: https://www.agriculture.go.ug.",
    "[10] World Health Organization, \"Aflatoxins,\" WHO Food Safety Digest, 2023.",
    "[11] International Agency for Research on Cancer, \"Aflatoxins,\" IARC Monographs, 2012.",
    "[12] United Nations, \"Transforming our world: The 2030 Agenda for Sustainable Development,\" New York: UN, 2015.",
    "[13] \"flutter_tts,\" Dart package, pub.dev. [Online]. Available: https://pub.dev/packages/flutter_tts.",
    "[14] \"workmanager,\" Dart package, pub.dev. [Online]. Available: https://pub.dev/packages/workmanager.",
]
for r in refs:
    doc.add_paragraph(r)

doc.add_page_break()

# ---- 2 Overall Description --------------------------------------------------
add_banner("2  Overall Description")

add_heading("2.1  Product Perspective", level=2)
add_para(
    "AflAlert is a new, self-contained mobile product, not a follow-on to an existing "
    "system. It is a cross-platform Flutter application (Android and iOS, with a "
    "limited Firebase-hosted web build) backed by Firebase for authentication, data "
    "storage, and server-side logic. Kernel classification runs entirely on-device via "
    "a bundled TensorFlow Lite model, so Tier 1 scanning works without a network "
    "connection; test-strip reading is also on-device, using deterministic image "
    "processing rather than a trained model. Account sync, scan-history storage, PDF "
    "report generation/sharing, and weather-based alerts depend on connectivity to "
    "Firebase and the Open-Meteo weather API."
)
if os.path.exists(ARCH_DIAGRAM_PATH):
    arch_p = doc.add_paragraph()
    arch_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    arch_p.add_run().add_picture(ARCH_DIAGRAM_PATH, width=content_width)
    add_para(
        "Figure 1: AflAlert system context — Flutter UI screens call into a services "
        "layer (classification, strip analysis, alerts, voice, PDF/report storage, "
        "Firebase access), which talks to the platform and external services below it.",
        italic=True, size=10,
    )

add_heading("2.2  Product Functionality", level=2)
add_bullets([
    "Kernel scan (Tier 1) — photograph maize kernels; an on-device TFLite model flags visible mould/discoloration risk.",
    "Test strip scan (Tier 2) — photograph a reacted lateral-flow strip; the app reads the Control/Test lines and estimates a ppb contamination reading.",
    "Guided recommendations — every result comes with UNBS/MAAIF-sourced next-step guidance.",
    "Weather and rain alerts — a daily morning summary plus urgent incoming-rain alerts so drying grain can be protected.",
    "Voice assistant — hands-free scanning and Q&A, reachable from the home, history, notifications, and result screens.",
    "Scan history and PDF reports — review past scans and export/share a PDF report per result.",
    "Accounts — email/Google sign-in, OTP verification, password reset, biometric app unlock.",
    "Localization — full English and Luganda support throughout the app.",
])

add_heading("2.3  Users and Characteristics", level=2)
add_lettered([
    ("Farmers and grain handlers (primary)", "Vary widely in literacy, smartphone experience, and connectivity. Use kernel scanning most, benefit most from the voice assistant and Luganda localization, and are the group the app's core UX decisions are optimised for."),
    ("Grain traders and buyers (secondary)", "Use scan results and PDF reports as a quick, shareable quality signal before a purchase or sale; more likely to also use Tier 2 strip scanning where a more quantitative reading matters."),
    ("Extension officers and inspectors (tertiary)", "Consume shared PDF reports and guidance content rather than performing scans themselves; interact with the app mainly at the point results are shared with them."),
    ("Group 4 (developers/administrators)", "Maintain the Firebase backend, Cloud Functions, and the bundled TFLite model; the only users with access to backend configuration, not represented in the end-user use cases in §3.3."),
])
add_para(
    "Farmers and grain handlers are the most important user group to satisfy: the "
    "kernel-scan tier is deliberately offline-capable and camera-only so it works for "
    "them regardless of connectivity or access to a test strip."
)

add_heading("2.4  Operating Environment", level=2)
add_para(
    "Client: Android (minSdk 24, i.e. Android 7.0+) and iOS, built with Flutter (Dart "
    "SDK ^3.12.2). A Firebase Hosting web build (https://aflalert.web.app) supports "
    "onboarding, authentication, history, and alerts, but not camera-based scanning, "
    "since the bundled TFLite classifier relies on dart:ffi native bindings with no "
    "browser-compatible implementation."
)
add_para(
    "Backend: Firebase Authentication, Cloud Firestore, Firebase Storage, Cloud "
    "Functions, and Firebase App Check, all hosted on Google Cloud. Weather data is "
    "fetched from the Open-Meteo forecast API. Background alert scheduling on Android "
    "uses WorkManager so alerts can fire while the app is closed."
)

add_heading("2.5  Design and Implementation Constraints", level=2)
add_lettered([
    ("On-device ML via dart:ffi", "The kernel classifier is a bundled TFLite model accessed through tflite_flutter/dart:ffi, which has no browser-compatible implementation — camera-based scanning is therefore Android/iOS-only, not available on the web build."),
    ("Firebase platform dependency", "Authentication, data storage, file storage, and backend logic are built directly on Firebase's Auth/Firestore/Storage/Functions/App Check services, which constrains hosting choice and offline-sync behaviour to what those SDKs provide."),
    ("Android-only background scheduling", "WorkManager provides no equivalent timing guarantee on iOS, so proactive rain/weather alerts currently only run reliably on Android."),
    ("Unvalidated ppb calibration", "The strip reader's ppb figure is computed from a monotonic placeholder curve (not a lab-fitted dose-response curve), constraining Tier 2 results to a relative/qualitative signal until certified reference strips are used to calibrate it."),
    ("Cross-platform language/framework", "The team committed to Flutter/Dart for a single codebase across Android, iOS, and web, and to maintaining full English/Luganda parity (app_localizations ARB files) for every user-facing string."),
    ("Course timeline and team size", "Delivered by a 5-person team within the CSC 1304 Practical Skills Development timeline, which bounds how much of the roadmap in §3.2 and the technical report's \"Next Steps\" can be completed before submission."),
])

add_heading("2.6  User Documentation", level=2)
add_para(
    "AflAlert relies on in-app guidance rather than a separate manual: an onboarding "
    "carousel introduces the app's purpose at first launch, and guided capture overlays "
    "walk the user through kernel and strip photo capture at the point of use. A short "
    "online FAQ/user guide is planned alongside the GitHub repository for users who "
    "want more detail than the in-app guidance provides. All help content is localized "
    "in English and Luganda, consistent with the rest of the app."
)

add_heading("2.7  Assumptions and Dependencies", level=2)
add_bullets([
    "Users have an Android or iOS smartphone with a working rear camera.",
    "Kernel scanning (Tier 1) works offline, but account sync, history storage, weather alerts, and PDF sharing assume at least intermittent internet connectivity.",
    "Firebase's free/Blaze-tier services remain available and sufficient for the project's usage volume.",
    "Third-party packages the app depends on (tflite_flutter, image, workmanager, flutter_tts, geolocator) remain compatible with the Dart SDK ^3.12.2 constraint.",
    "The Open-Meteo forecast API remains accessible and reasonably accurate for the Ugandan districts the app targets.",
    "Users performing a Tier 2 scan have access to a commercially available lateral-flow aflatoxin test strip and can follow the extraction step correctly.",
    "The Tier 2 ppb reading is assumed to be a relative/qualitative signal, not a certified quantitative assay result, until it is calibrated against certified reference standards.",
])

doc.add_page_break()

# ---- 3 Specific Requirements ------------------------------------------------
add_banner("3  Specific Requirements")

add_heading("3.1  External Interface Requirements", level=2)

add_heading("3.1.1  User Interfaces", level=3)
add_para(
    "AflAlert's screens fall into four groups: onboarding/authentication (onboarding "
    "carousel, welcome, login, registration, OTP verification, forgot-password), the "
    "home screen (scan launch buttons, recent scans, weather/alert banner), the two "
    "scan flows (guided capture → result, for both kernel and strip), and supporting "
    "screens (history, notifications, voice assistant overlay, settings, profile). "
    "Every screen exposes a consistent back/cancel action; capture screens show an "
    "in-frame guide overlay; result screens use a consistent safe/unsafe colour and "
    "iconography; invalid or low-confidence scans are shown as an inline error/retry "
    "banner rather than a silent failure or a misleading result."
)

add_heading("3.1.2  Hardware Interfaces", level=3)
add_bullets([
    "Rear camera — used for kernel and test-strip photo capture (camera / image_picker).",
    "Location sensor — device GPS/network location (geolocator, geocoding) used to resolve the user's district for weather and rain alerts.",
    "Biometric hardware — fingerprint/face unlock (local_auth) used for optional app-unlock on a signed-in session.",
    "Touchscreen — the sole input mechanism outside of the voice assistant's microphone input.",
])

add_heading("3.1.3  Software Interfaces", level=3)
add_bullets([
    "Android and iOS platform APIs (camera, notifications, biometrics, background execution).",
    "Firebase SDKs — Authentication, Cloud Firestore, Storage, Cloud Functions, App Check.",
    "TensorFlow Lite runtime (tflite_flutter) for on-device kernel classification.",
    "WorkManager (Android) for background rain/weather alert scheduling.",
    "flutter_local_notifications for displaying alerts and scan-complete notifications.",
])

add_heading("3.1.4  Communications Interfaces", level=3)
add_para(
    "All communication with the Firebase backend and Cloud Functions is over HTTPS, "
    "authenticated with a Firebase Auth session token and validated with a Firebase App "
    "Check token to reject requests not originating from the genuine app. Weather data "
    "is fetched over HTTPS from the Open-Meteo REST API using the device's resolved "
    "coordinates. Push/local notifications carry the daily weather summary and urgent "
    "rain alerts. PDF reports are handed off to the platform share sheet (e-mail, "
    "WhatsApp, etc.) rather than transmitted by the app itself. No custom encryption is "
    "implemented beyond standard TLS in transit and the security controls already "
    "provided by Firebase."
)

add_heading("3.2  Functional Requirements", level=2)

add_heading("3.2.1  Kernel Image Scanning", level=3)
add_req_table([
    ("FR-KER-1", "The user shall be able to capture a photo of a maize kernel sample via the in-app camera, or select an existing photo from the gallery."),
    ("FR-KER-2", "The system shall classify the captured/selected image on-device via the bundled TFLite model (e.g. healthy vs. mouldy/spoiled) without requiring a network connection."),
    ("FR-KER-3", "The system shall reject low-confidence or clearly non-maize images (colour-profile mismatch, model rejection, low confidence) instead of returning a result."),
])

add_heading("3.2.2  Test Strip Scanning", level=3)
add_req_table([
    ("FR-STR-1", "The system shall present a guided capture frame aligning the Control line near the top and Test line near the bottom, and check light/focus quality before capture."),
    ("FR-STR-2", "The system shall reject photos that do not plausibly resemble a strip cassette before attempting line detection."),
    ("FR-STR-3", "The system shall locate the Control (C) and Test (T) line bands from a row-by-row luminance profile, using a local-background baseline for robustness to uneven lighting."),
    ("FR-STR-4", "The system shall compute the optical density of each line relative to its local background, derive a Test/Control ratio, and translate it into an estimated ppb reading, classified safe when at or below 20 ppb."),
    ("FR-STR-5", "The system shall flag a scan as invalid (rather than presenting a misleading low-ppb result) if no strip/lines are detected, or if the Control line is not dark enough to indicate a properly developed run."),
])

add_heading("3.2.3  Guided Recommendations", level=3)
add_req_table([
    ("FR-REC-1", "The system shall attach next-step guidance, sourced from UNBS and MAAIF standards, to every scan result (kernel or strip), based on whether the result is within or above the applicable safety threshold."),
])

add_heading("3.2.4  Weather and Rain Alerts", level=3)
add_req_table([
    ("FR-ALT-1", "The system shall provide a once-daily weather summary based on the user's resolved location."),
    ("FR-ALT-2", "The system shall independently monitor short-range forecasts and push an urgent alert when rain is imminent, including while the app is closed (Android background scheduling)."),
])

add_heading("3.2.5  Voice Assistant", level=3)
add_req_table([
    ("FR-VOI-1", "The user shall be able to trigger scans and ask questions about results/recommendations using speech input, with spoken responses."),
    ("FR-VOI-2", "The assistant shall be reachable from the home, history, notifications, and both result screens."),
    ("FR-VOI-3", "When a scan is triggered from the assistant, the system shall speak a summary of the result (safe/unsafe) via text-to-speech as soon as it is ready, using the same label logic as the on-screen result."),
])

add_heading("3.2.6  History, Reports and Accounts", level=3)
add_req_table([
    ("FR-HIS-1", "The system shall store and display a history of past kernel and strip scans per authenticated user."),
    ("FR-HIS-2", "The user shall be able to generate and share/download a PDF report for a given scan result."),
    ("FR-ACC-1", "The system shall support email/password and Google sign-in, OTP verification, password reset, and biometric app unlock."),
    ("FR-ACC-2", "The system shall support English and Luganda throughout the app, switchable from Settings."),
])

add_heading("3.3  Behaviour Requirements", level=2)
add_heading("3.3.1  Use Case View", level=3)
add_para(
    "The primary actor is the Farmer / Grain Handler, who drives every use case except "
    "the alert-triggering itself. The Background Scheduler (Android WorkManager) is a "
    "secondary actor that independently triggers the “Receive Weather & Rain Alerts” "
    "use case even while the app is closed. Use cases:"
)
add_bullets([
    "Register / Sign In — create an account or sign in (email/password or Google), including OTP verification.",
    "Scan Maize Kernels (Tier 1) — capture/select a kernel photo and get an on-device risk classification.",
    "Scan Test Strip (Tier 2) — capture a reacted strip photo and get a ppb-based risk reading.",
    "View Guided Recommendations — view UNBS/MAAIF-sourced next steps attached to a scan result.",
    "Receive Weather & Rain Alerts — receive the daily summary and urgent rain notifications.",
    "Use Voice Assistant — trigger scans and ask result/recommendation questions hands-free.",
    "View Scan History — browse past kernel and strip scan results.",
    "Export PDF Report — generate and share a PDF of a scan result.",
    "Manage Settings & Profile — change language, notification preferences, and account details.",
])
if os.path.exists(USE_CASE_DIAGRAM_PATH):
    uc_p = doc.add_paragraph()
    uc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uc_p.add_run().add_picture(USE_CASE_DIAGRAM_PATH, width=content_width)
    add_para(
        "Figure 2: AflAlert use case diagram — Farmer/Grain Handler as primary actor "
        "against all end-user use cases; Background Scheduler as secondary actor "
        "triggering rain/weather alerts independently of the app being open.",
        italic=True, size=10,
    )

doc.add_page_break()

# ---- 4 Other Non-functional Requirements ------------------------------------
add_banner("4  Other Non-functional Requirements")

add_heading("4.1  Performance Requirements", level=2)
add_numbered([
    "Kernel scan classification shall return a result within 3 seconds of capture on a mid-range device, since inference runs on-device with no network round-trip.",
    "Test-strip line detection and ppb calculation shall complete within 5 seconds of capture.",
    "The app shall reach the home screen within 4 seconds of a cold start on a mid-range Android device.",
    "A scan-history list of up to 200 entries shall load within 2 seconds from Firestore under normal connectivity.",
    "PDF report generation and hand-off to the platform share sheet shall complete within 3 seconds of the user's request.",
    "The daily weather summary notification shall be delivered within a fixed morning window (e.g. 6:00–8:00 AM local time) on at least 95% of scheduled days.",
])

add_heading("4.2  Safety and Security Requirements", level=2)
add_heading("4.2.1  Safety", level=3)
add_numbered([
    "Every result screen (kernel or strip) shall display a disclaimer that the reading is a first-line risk indicator, not a certified laboratory diagnosis, given the seriousness of chronic aflatoxin exposure.",
    "Results above the applicable safety threshold shall carry explicit “do not mix with clean grain” / “set aside for inspection” guidance, to prevent contaminated grain re-entering the food supply.",
    "Invalid or low-confidence scans (poor lighting, undetected lines, non-maize image) shall be rejected rather than silently returned as a “safe” result.",
    "High-risk (above-threshold) results shall recommend confirmatory laboratory testing rather than being presented as a final determination.",
])
add_heading("4.2.2  Security", level=3)
add_bullets([
    "Scan history and reports are stored per-account and require the user to be authenticated (email/password or Google sign-in).",
    "Registration and password reset require OTP verification.",
    "Biometric app-unlock (local_auth) is available to protect an active session on a shared device.",
    "Firebase App Check verifies that backend requests originate from the genuine AflAlert app.",
    "All network traffic is transmitted over HTTPS/TLS.",
    "Firestore security rules restrict each user to reading and writing only their own scan history and profile data.",
    "No scan photos or location data are shared with third parties beyond what is required for weather lookups.",
])

add_heading("4.3  Software Quality Attributes", level=2)
add_heading("4.3.1  Reliability", level=3)
add_para(
    "Kernel scanning (Tier 1) works fully offline, so the app's core function is not "
    "affected by connectivity loss. Both scan tiers detect and reject invalid input "
    "(low confidence, undetected lines, poor-quality capture) rather than failing "
    "silently or returning a misleading result."
)
add_heading("4.3.2  Portability", level=3)
add_para(
    "A single Flutter/Dart codebase targets Android, iOS, and a limited Firebase-hosted "
    "web build, sharing UI, services, and localization resources across all three."
)
add_heading("4.3.3  Usability", level=3)
add_para(
    "Guided capture overlays, a hands-free voice assistant, and full English/Luganda "
    "localization are prioritised specifically because the primary users — smallholder "
    "farmers and grain handlers — vary widely in literacy and smartphone experience. "
    "Ease of use is prioritised over feature breadth for this group."
)
add_heading("4.3.4  Maintainability", level=3)
add_para(
    "The codebase separates screens (UI), services (business logic — classification, "
    "strip analysis, Firebase access, alerts, PDF generation, voice), and localization "
    "resources, so each detection tier and supporting feature can be modified and "
    "tested independently."
)
add_heading("4.3.5  Availability", level=3)
add_para(
    "Backend-dependent features (sync, history, alerts, PDF sharing) inherit Firebase's "
    "published service availability; the core kernel-scan feature remains available "
    "even during a backend outage, since it does not depend on the network."
)

doc.add_page_break()

# ---- 5 Other Requirements ---------------------------------------------------
add_banner("5  Other Requirements")
add_lettered([
    ("Localization", "Full English and Luganda parity is required across all user-facing strings (app_localizations ARB files); the localization architecture should allow additional languages to be added without changing screen code."),
    ("Legal / data handling", "A Terms of Service and data-handling disclosure covering photo and location storage shall be presented at sign-up."),
    ("Data retention", "Scan photos and history are retained per-account until the user deletes their account, at which point associated Firestore and Storage data shall be removed."),
])

doc.add_page_break()

# ---- Appendix B - Group Log --------------------------------------------------
add_banner("Appendix B - Group Log", toc_level=None)
add_placeholder_para(
    "TO DO: Log the actual Group 4 meetings here (date, attendees, decisions made) "
    "and any meetings with users/the mentor, before final submission. This section is "
    "intentionally left for the group to fill in with real minutes rather than "
    "placeholder content, per the course facilitator's requirement to demonstrate "
    "effort put into producing this document."
)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
for i, h in enumerate(["Date", "Attendees", "Summary / Decisions"]):
    hdr[i].text = h
    shade_cell(hdr[i], BANNER_FILL)
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = BANNER_TEXT
row = table.add_row().cells
for c in row:
    p = c.paragraphs[0]
    r = p.add_run("<to be filled in>")
    r.italic = True
    r.font.color.rgb = PLACEHOLDER

# ---- Save ---------------------------------------------------------------
os.makedirs(DESKTOP_DIR, exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Saved {OUTPUT_PATH}")
